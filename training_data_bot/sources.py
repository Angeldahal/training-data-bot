from abc import ABC, abstractmethod
import asyncio
import csv
import datetime
import json
from pathlib import Path
from typing import List
from uuid import uuid4

from .core.logging import get_logger
from .core.exceptions import DocumentLoadError
from .models import Document, DocumentType


class BaseLoader(ABC):
    def __init__(self):
        self.logger = get_logger(f"loader.{self.__class__.__name__}")
        self.supported_formats: List[DocumentType] = []

    @abstractmethod
    async def load_single(self, source, **kwargs) -> Document:
        pass

    async def load_multiple(self, sources, max_workers=4):
        semaphore = asyncio.Semaphore(max_workers)
        async def load_with_semaphore(source):
            async with semaphore:
                return await self.load_single(source)
        tasks = [load_with_semaphore(source) for source in sources]
        results = await asyncio.gather(*tasks)
    
    def get_document_type(self, source):
        if source.startswith("http"):
            return DocumentType.URL
        source = Path(source)
        suffix = source.suffix.lower().lstrip(".")
        return DocumentType(suffix)

    def create_document(
            self,
            title,
            content,
            source,
            doc_type,
            **kwargs
    ):
        return Document(
            id=uuid4(),
            title=title,
            content=content,
            source=source,
            doc_type=doc_type,
            word_count=len(content.split()),
            created_at=datetime.utcnow(),
            **kwargs
        )


class PDFLoader:
    pass


class WebLoader:
    pass


class DocumentLoader(BaseLoader):
    def __init__(self):
        super().__init__()
        self.supported_formats = [
            DocumentType.TXT,
            DocumentType.MD,
            DocumentType.JSON,
            DocumentType.CSV,
            DocumentType.HTML,
            DocumentType.DOCX,
        ]

    async def load_single(self, source, encoding='utf-8'):
        doc_type = self.get_document_type(source)
        if doc_type == DocumentType.TXT:
            content = await self._load_text(source, encoding)
        elif doc_type == DocumentType.MD:
            content = await self._load_markdown(source, encoding)
        elif doc_type == DocumentType.HTML:
            content = await self._load_html(source, encoding)
        elif doc_type == DocumentType.CSV:
            content = await self._load_csv(source, encoding)
        elif doc_type == DocumentType.JSON:
            content = await self._load_json(source, encoding)
        elif doc_type == DocumentType.DOCX:
            content = await self._load_docx(source)
    
    async def _load_text(self, path, encoding):
        return await asyncio.to_thread(path.read_text, encoding=encoding)

    async def _load_md(self, path, encoding):
        return await asyncio.to_thread(path.read_text, encoding=encoding)
    
    async def _load_html(self, path, encoding):
        try:
            from bs4 import BeautifulSoup
            with open(path, "r", encoding=encoding) as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split(" "))
            return " ".join(chunk for chunk in chunks if chunk)
        except ImportError:
            return path.read_text(encoding=encoding)

    async def _load_json(self, path, encoding):
        with open(path, "r", encoding=encoding) as f:
            data = json.load(f)
        if isinstance(data, dict):
            lines = [f"{key}: {value}" for key, value in data.items()]
            return '\n'.join(lines)
        elif isinstance(data, list):
            lines = [f"Item {i+1}: {item}" for i, item in enumerate(data)]
            return "\n".join(lines)
    
    async def _load_csv(self, path, encoding):
        lines = []
        with open(path, 'r', encoding=encoding, newline='') as f:
            reader = csv.reader(f)
            headers = next(reader, None)
            if headers:
                lines.append("Headers: " + ", ".join(headers))
                lines.append("")
            for row_num, row in enumerate(reader, 1):
                if headers and len(row) == len(headers):
                    row_data = [f'{header}: {value}' for header, value in zip(headers, row)]
                    lines.append(f'Row: {row_num}: {' | '.join(row_data)}')
        return "\n".join(lines)

    async def _load_docx(self, path, encoding):
        try:
            from docx import Document
            doc = Document(path)
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(text_parts)
        except ImportError:
            raise DocumentLoadError("python-docx package required for DOCX files")


class UnifiedLoader(BaseLoader):
    def __init__(self):
        self.document_loader = DocumentLoader()
        self.pdf_loader = PDFLoader()
        self.web_loader = WebLoader()

        self.supported_formats = list(DocumentType)

    def _get_loader(self, source):
        if source.startswith(("http://", "https://")):
            return self.web_loader

        source = Path(source)
        if not source.exists():
            return None

        suffix = source.suffix.lower().lstrip(".")
        doctype = DocumentType(suffix)

        if doctype == DocumentType.PDF:
            return self.pdf_loader
        elif doctype in [DocumentType.TXT, DocumentType.DOCX, DocumentType.MD]:
            return self.document_loader

    def load_single(self):
        pass

    def load_directory(self):
        pass
